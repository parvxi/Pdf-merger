"""
DCL PDF Merger API — v5.0.0 (faithful to v4.0.0 + batched conversion)

This is your ORIGINAL v4.0.0 with ONE behavioural change: instead of launching a
fresh LibreOffice process per Office file (checklist + each .docx/.xls/.xlsx),
ALL Office files are converted in a SINGLE LibreOffice invocation. That removes
the per-file cold-start cost that pushed big merges past Power Automate's ~120s
synchronous ceiling (the 502 "NoResponse").

PRESERVED EXACTLY (do not regress these — the flow depends on them):
  * Request model: files[], output_name, templateData (+ legacy checklist fields)
  * Content-first detect_file_type (pdf/png/jpg/gif/webp/docx/xlsx/doc/xls)
  * Checklist generation from templates/DCL_Template.docx via replace_placeholders
  * Response shape: { success, output_name, size_bytes, content(base64), files_merged, conversions }
    -> the flow reads the merged PDF from `content`. NEVER rename this field.

ADDED:
  * Batched Office->PDF conversion (one LibreOffice call) with per-file fallback
  * Per-file append guard + `skipped` list so one bad file can't 500 the package
"""

from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import List, Dict
import base64
from io import BytesIO
from pypdf import PdfWriter
import logging
import openpyxl
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
import tempfile
import os
import subprocess
import shutil
import uuid
from PIL import Image
from docx import Document

# ==========================================
# Logging
# ==========================================
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# ==========================================
# Check LibreOffice
# ==========================================
path = shutil.which("libreoffice")
if path:
    logger.info(f"✅ LibreOffice found at: {path}")
else:
    logger.warning("⚠ LibreOffice NOT FOUND — DOC/DOCX/XLS conversion will fail.")

# A single, reused LibreOffice profile dir avoids the ~3-5s first-run bootstrap.
LO_PROFILE = os.path.join(tempfile.gettempdir(), "lo_profile")

# One batched conversion of all Office files. Keep under the flow's 120s ceiling.
LO_TIMEOUT_SECONDS = 110

# ==========================================
# FastAPI App
# ==========================================
app = FastAPI(
    title="DCL PDF Merger API",
    description="Convert ANY uploaded DCL files to PDF and merge",
    version="5.0.0",
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
    allow_credentials=True,
)


# ==========================================
# Pydantic Models  (UNCHANGED — the flow's contract)
# ==========================================
class DocumentFile(BaseModel):
    name: str
    content: str  # base64


class MergeRequest(BaseModel):
    files: List[DocumentFile]
    output_name: str = "merged_dcl.pdf"

    # NEW STRUCTURE - Matches extraction system
    templateData: dict | None = None  # All 67 fields in one object

    # DEPRECATED - Keep for backward compatibility only
    checklist: dict | None = None
    checklistMapped: dict | None = None
    templateMaster: dict | None = None


# ==========================================
# DOCX Placeholder Replacement  (UNCHANGED)
# ==========================================
def replace_placeholders(doc, data):

    def replace_in_paragraph(p):
        full_text = "".join(run.text for run in p.runs)
        for key, value in data.items():
            full_text = full_text.replace("{{" + key + "}}", str(value))
        for run in p.runs:
            run.text = ""
        if p.runs:
            p.runs[0].text = full_text
        else:
            p.add_run(full_text)

    def replace_in_table(table):
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_paragraph(p)
                for t2 in cell.tables:
                    replace_in_table(t2)

    for p in doc.paragraphs:
        replace_in_paragraph(p)
    for t in doc.tables:
        replace_in_table(t)
    for section in doc.sections:
        footer = section.footer
        for p in footer.paragraphs:
            replace_in_paragraph(p)
        for t in footer.tables:
            replace_in_table(t)


def build_checklist_docx(full_data: dict, workdir: str) -> str:
    """
    Fill templates/DCL_Template.docx from `full_data` and save a .docx into
    workdir. Returns the docx path so it can be converted in the SAME batch as
    the other Office files (instead of its own LibreOffice call). The checklist
    content/logic is identical to the original generate_checklist_pdf.
    """
    template_path = "templates/DCL_Template.docx"
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template missing: {template_path}")

    doc = Document(template_path)
    replace_placeholders(doc, full_data)

    out = os.path.join(workdir, f"checklist_{uuid.uuid4().hex}.docx")
    doc.save(out)
    return out


@app.get("/debug-template")
def debug_template():
    path = "templates/DCL_Template.docx"
    return {
        "exists": os.path.exists(path),
        "absolute_path": os.path.abspath(path),
        "cwd": os.getcwd(),
        "dir_listing": os.listdir("templates") if os.path.exists("templates") else "templates folder missing",
    }


# ==========================================
# File Type Detection  (UNCHANGED — content-first, correct)
# ==========================================
def detect_file_type(content_bytes: bytes, filename: str) -> str:
    if content_bytes.startswith(b"%PDF"):
        return "pdf"
    if content_bytes.startswith(b"\x89PNG\r\n\x1a\n"):
        return "png"
    if content_bytes[0:3] == b"\xff\xd8\xff":
        return "jpg"
    if content_bytes.startswith(b"GIF87a") or content_bytes.startswith(b"GIF89a"):
        return "gif"
    if content_bytes.startswith(b"RIFF") and content_bytes[8:12] == b"WEBP":
        return "webp"
    if content_bytes.startswith(b"PK\x03\x04"):
        header = content_bytes[:2000]
        if b"word/" in header:
            return "docx"
        if b"xl/" in header:
            return "xlsx"
    if content_bytes.startswith(b"\xd0\xcf\x11\xe0"):
        if filename.lower().endswith(".doc"):
            return "doc"
        if filename.lower().endswith(".xls"):
            return "xls"
    ext = filename.lower().split(".")[-1]
    if ext in ["pdf", "docx", "xlsx", "doc", "xls", "png", "jpg", "jpeg", "gif", "webp"]:
        return ext
    return "unknown"


# ==========================================
# Image -> PDF  (UNCHANGED, in-process, no LibreOffice)
# ==========================================
def convert_image_to_pdf(image_bytes: bytes) -> bytes:
    try:
        img = Image.open(BytesIO(image_bytes))
        if img.mode in ("RGBA", "LA"):
            bg = Image.new("RGB", img.size, (255, 255, 255))
            bg.paste(img, mask=img.split()[-1])
            img = bg
        else:
            img = img.convert("RGB")
        output = BytesIO()
        img.save(output, format="PDF")
        return output.getvalue()
    except Exception as e:
        raise ValueError(f"Failed to convert image: {e}")


# ==========================================
# Office -> PDF  (NEW: one LibreOffice call for ALL office files)
# ==========================================
def convert_office_files_batch(paths: List[str], outdir: str) -> Dict[str, str]:
    """
    Convert every given Office file (.docx/.doc/.xlsx/.xls) to PDF in a SINGLE
    LibreOffice invocation. Returns {source_path: produced_pdf_path}. Anything
    the batch misses is retried individually so one bad file can't sink the rest.
    """
    if not paths:
        return {}

    os.makedirs(LO_PROFILE, exist_ok=True)
    cmd = [
        "libreoffice",
        "-env:UserInstallation=file://" + LO_PROFILE,
        "--headless", "--norestore", "--nolockcheck", "--nodefault",
        "--convert-to", "pdf",
        "--outdir", outdir,
        *paths,
    ]
    result_map: Dict[str, str] = {}
    try:
        subprocess.run(
            cmd, check=True, timeout=LO_TIMEOUT_SECONDS,
            stdout=subprocess.PIPE, stderr=subprocess.PIPE,
        )
        for src in paths:
            produced = os.path.join(outdir, os.path.splitext(os.path.basename(src))[0] + ".pdf")
            if os.path.exists(produced):
                result_map[src] = produced
    except (subprocess.CalledProcessError, subprocess.TimeoutExpired) as exc:
        logger.warning(f"[batch-convert] failed ({exc}); falling back to per-file")

    for src in [p for p in paths if p not in result_map]:
        produced = _convert_one(src, outdir)
        if produced:
            result_map[src] = produced
    return result_map


def _convert_one(src: str, outdir: str) -> str:
    os.makedirs(LO_PROFILE, exist_ok=True)
    cmd = [
        "libreoffice",
        "-env:UserInstallation=file://" + LO_PROFILE,
        "--headless", "--norestore", "--nolockcheck", "--nodefault",
        "--convert-to", "pdf", "--outdir", outdir, src,
    ]
    try:
        subprocess.run(
            cmd, check=True, timeout=LO_TIMEOUT_SECONDS,
            stdout=subprocess.PIPE, stderr=subprocess.PIPE,
        )
        produced = os.path.join(outdir, os.path.splitext(os.path.basename(src))[0] + ".pdf")
        return produced if os.path.exists(produced) else ""
    except Exception as exc:  # noqa: BLE001
        logger.warning(f"[convert-one] {src} failed: {exc}")
        return ""


# ==========================================
# Merge Endpoint
# ==========================================
@app.post("/merge-pdfs")
async def merge_pdfs(request: MergeRequest):
    if not request.files:
        raise HTTPException(400, "No files provided.")

    # Resolve checklist data (new templateData or legacy structure) — UNCHANGED
    if request.templateData:
        full_checklist_data = request.templateData
        logger.info("✅ Using NEW templateData structure")
    else:
        full_checklist_data = {
            **(request.checklist or {}),
            **(request.checklistMapped or {}),
            **(request.templateMaster or {}),
        }
        logger.info("⚠️ Using OLD structure (checklist + checklistMapped + templateMaster)")

    workdir = tempfile.mkdtemp(prefix="merge_")
    conversions: Dict[str, str] = {}
    try:
        # ordered[]: merge order preserved. office items are converted in batch.
        ordered: List[dict] = []
        office_srcs: List[str] = []

        # 1) Checklist first (as an Office docx folded into the batch)
        if full_checklist_data:
            try:
                checklist_src = build_checklist_docx(full_checklist_data, workdir)
                ordered.append({"name": "DCL Checklist", "kind": "office", "src": checklist_src})
                office_srcs.append(checklist_src)
            except Exception as e:
                raise HTTPException(500, f"Checklist PDF generation failed: {e}")

        # 2) Classify each uploaded file -> pdf passthrough / image / office
        for idx, f in enumerate(request.files):
            try:
                file_bytes = base64.b64decode(f.content)
            except Exception:
                raise HTTPException(400, f"Invalid base64 for {f.name}")

            ftype = detect_file_type(file_bytes, f.name)

            if ftype == "pdf":
                p = os.path.join(workdir, f"{idx:03d}_{uuid.uuid4().hex}.pdf")
                with open(p, "wb") as fh:
                    fh.write(file_bytes)
                ordered.append({"name": f.name, "kind": "pdf", "pdf": p})
                conversions[f.name] = "passed_pdf"

            elif ftype in ("png", "jpg", "jpeg", "gif", "webp"):
                pdf_bytes = convert_image_to_pdf(file_bytes)
                p = os.path.join(workdir, f"{idx:03d}_{uuid.uuid4().hex}.pdf")
                with open(p, "wb") as fh:
                    fh.write(pdf_bytes)
                ordered.append({"name": f.name, "kind": "pdf", "pdf": p})
                conversions[f.name] = "image_to_pdf"

            elif ftype in ("docx", "doc", "xlsx", "xls"):
                src = os.path.join(workdir, f"{idx:03d}_{uuid.uuid4().hex}.{ftype}")
                with open(src, "wb") as fh:
                    fh.write(file_bytes)
                ordered.append({"name": f.name, "kind": "office", "src": src})
                office_srcs.append(src)
                conversions[f.name] = f"{ftype}_to_pdf_batch"

            else:
                # Robust: don't 500 the whole package over one odd attachment.
                conversions[f.name] = "unsupported_skipped"
                logger.warning(f"[merge] unsupported file type, skipped: {f.name}")

        # 3) ONE batched LibreOffice conversion for all Office files
        converted = convert_office_files_batch(office_srcs, workdir)

        # 4) Merge in order, guarding each append
        merger = PdfWriter()
        skipped: List[str] = []
        for item in ordered:
            pdf_path = item["pdf"] if item["kind"] == "pdf" else converted.get(item["src"], "")
            if not (pdf_path and os.path.exists(pdf_path)):
                skipped.append(item["name"])
                logger.warning(f"[merge] no PDF produced for '{item['name']}' — skipped")
                continue
            try:
                merger.append(pdf_path)
            except Exception as exc:  # noqa: BLE001
                skipped.append(item["name"])
                logger.warning(f"[merge] could not append '{item['name']}' ({exc}) — skipped")

        if len(merger.pages) == 0:
            raise HTTPException(500, "No pages to merge — every input failed: " + ", ".join(skipped))

        output_stream = BytesIO()
        merger.write(output_stream)
        merger.close()
        final_pdf = output_stream.getvalue()

        if skipped:
            logger.warning(f"[merge] completed with {len(skipped)} skipped: {skipped}")

        # Response shape UNCHANGED (+ skipped). The flow reads `content`.
        return {
            "success": True,
            "output_name": request.output_name,
            "size_bytes": len(final_pdf),
            "content": base64.b64encode(final_pdf).decode(),
            "files_merged": len(request.files),
            "conversions": conversions,
            "skipped": skipped,
        }
    finally:
        shutil.rmtree(workdir, ignore_errors=True)


@app.get("/health")
def health():
    return {"status": "ok"}


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
